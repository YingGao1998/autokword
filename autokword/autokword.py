import pandas as pd
from docx import Document
from docx.shared import Pt
import os

# 所有文件已经上传到Google Colab，并在/content/目录下
base_dir = '/content/'

# 定义需要处理的 sheet 名称
sheets = ['CA0', 'CA1', 'US0', 'US1']

# 创建一个新的 Word 文档
doc = Document()


# 定义一个函数来插入标题
def add_title(title, size=14):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(size)


# 定义一个函数来插入表格
def add_table(data, headers, add_category=False):
    table = doc.add_table(rows=1, cols=len(headers))

    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)

    # 添加数据行
    for index, row in data.iterrows():
        row_cells = table.add_row().cells

        # 确保每行的单元格数量与数据列数一致
        for i in range(len(row)):
            if i < len(row_cells):
                try:
                    if pd.isna(row.iloc[i]):
                        row_cells[i].text = 'N/A'
                    else:
                        row_cells[i].text = str(row.iloc[i])
                except Exception as e:
                    row_cells[i].text = 'Error'

        # 如果需要添加馆别信息
        if add_category:
            category_column = len(row)  # 新的一列用于显示馆别
            category = get_keyword_category(row['Keyword'])  # 获取馆别信息
            if category_column < len(row_cells):
                row_cells[category_column].text = category if category else '未匹配'

    # 定义一个函数来根据关键词匹配所属的馆


def get_keyword_category(keyword):
    for category, keywords_set in category_keywords.items():
        if keyword in keywords_set:
            return category
    return None


# 读取本地 Excel 表格
def read_local_excel(file_path):
    try:
        return pd.read_excel(file_path)
    except Exception as e:
        print(f"读取 Excel 文件时出错: {e}")
        return pd.DataFrame()


# 读取本地 Word 文件内容，提取所有的文本作为关键词
def read_local_word_file(file_path):
    try:
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            content.append(para.text.strip())  # 去除多余的空白
        print(f"关键词集合: {set(content)}")
        return set(content)  # 返回作为关键词的集合
    except Exception as e:
        print(f"读取 Word 文件时出错: {e}")
        return set()


# 使用 os.path.join 构建文件路径
category_files = {
    '母婴宠物馆': os.path.join(base_dir, '母婴宠物馆.docx'),
    '生活馆': os.path.join(base_dir, '生活馆.docx'),
    '美护馆': os.path.join(base_dir, '美护馆.docx'),
    '美食馆': os.path.join(base_dir, '美食馆.docx')
}

# 读取分类文件中的关键词
category_keywords = {}
for category, file_path in category_files.items():
    category_keywords[category] = read_local_word_file(file_path)

# 定义 brand 表路径，日仓和第三方 docx 文件路径
brand_file = os.path.join(base_dir, 'brand 2024-09-06.xlsx')
ricang_file = os.path.join(base_dir, '日仓.docx')
third_party_file = os.path.join(base_dir, '第三方.docx')

# 从本地读取 brand 表，日仓和第三方表通过 Word 文件读取关键词
brand_data = read_local_excel(brand_file)
ricang_keywords = read_local_word_file(ricang_file)  # 从日仓 Word 文档中读取关键词
third_party_keywords = read_local_word_file(third_party_file)  # 从第三方 Word 文档中读取关键词

# 9.9-9.15 和 9.2-9.8 文件路径
week1_file = os.path.join(base_dir, '9.9-9.15.xlsx')
week2_file = os.path.join(base_dir, '9.2-9.8.xlsx')

# 处理每个 sheet 数据
for sheet in sheets:
    # 添加区域标题
    add_title(f'区域: {sheet}', size=16)

    # 读取本周和上周的四个表
    try:
        week1_data = pd.read_excel(week1_file, sheet_name=sheet)  # 本周
        week2_data = pd.read_excel(week2_file, sheet_name=sheet)  # 上周
    except Exception as e:
        print(f"读取 Excel 数据时出错: {e}")
        continue

    # 重命名列，确保理解表格中的数据
    week1_data.columns = ['Keyword', 'SearchCount_Week1']
    week2_data.columns = ['Keyword', 'SearchCount_Week2']

    # 将 'SearchCount_Week2' 列转换为 object 类型以避免类型不兼容
    week2_data['SearchCount_Week2'] = week2_data['SearchCount_Week2'].astype(object)

    # 合并两周的数据，基于关键词合并
    merged_data = pd.merge(week1_data, week2_data, on='Keyword', how='outer')

    # 如果某个关键词在上周不存在，则填充为 N/A
    merged_data['SearchCount_Week2'].fillna('N/A', inplace=True)

    # 计算变化值，忽略 N/A 的情况
    merged_data['Change'] = merged_data.apply(
        lambda row: row['SearchCount_Week1'] - row['SearchCount_Week2'] if row['SearchCount_Week2'] != 'N/A' else 'New',
        axis=1
    )

    # 1. 本周搜索次数最多的关键词，取第一个关键词
    top_search_keyword = merged_data[['Keyword', 'SearchCount_Week1']].sort_values(by='SearchCount_Week1',
                                                                                   ascending=False).head(1)
    add_title('本周搜索次数最多的关键词', size=12)
    add_table(top_search_keyword, ['关键词', '搜索次数'])

    # 2. 新增次数较多的关键词：本周有，上周没有，且增加了40次以上的关键词
    new_keywords = merged_data[merged_data['SearchCount_Week2'] == 'N/A']
    new_keywords = new_keywords[new_keywords['SearchCount_Week1'] > 40]
    add_title('新增次数较多的关键词（增加了40次以上）', size=12)
    add_table(new_keywords, ['关键词', '搜索次数'])

    # 3. 单周搜索次数多的关键词：只取本周的数据，按前30排序
    single_week_keywords = week1_data[['Keyword', 'SearchCount_Week1']].sort_values(by='SearchCount_Week1',
                                                                                    ascending=False).head(30)
    add_title('单周搜索次数多的关键词（按前30排序）', size=12)
    add_table(single_week_keywords, ['关键词', '搜索次数', '分馆'], add_category=True)

    # 检查前30关键词是否出现在本地的四个表中：母婴宠物馆，生活馆，美护馆，美食馆
    unmatched_keywords = []
    for keyword in single_week_keywords['Keyword']:
        found = False
        for category, keywords_set in category_keywords.items():
            if keyword in keywords_set:
                found = True
                break
        if not found:
            unmatched_keywords.append(keyword)

    # 如果有未匹配的关键词，添加提示
    if unmatched_keywords:
        add_title('前30关键词中未匹配的关键词', size=12)
        for kw in unmatched_keywords:
            doc.add_paragraph(f"{kw} - 未匹配于任何分馆", style='List Bullet')

    # 4. 单周搜索次数增长多的关键词：先提取本周搜索次数前50个关键词，然后在其中筛选增长30%以上的关键词
    top_50_keywords = merged_data.sort_values(by='SearchCount_Week1', ascending=False).head(50)  # 本周前50个关键词

    # 筛选出前50个关键词中，搜索次数增长30%以上的关键词
    growth_keywords = top_50_keywords[(top_50_keywords['SearchCount_Week2'] != 'N/A') &
                                      ((top_50_keywords['SearchCount_Week1'] - top_50_keywords['SearchCount_Week2']) /
                                       top_50_keywords['SearchCount_Week2'] > 0.3)]

    # 显示筛选出的增长30%以上的关键词
    add_title('搜索次数增长多的关键词（增加30%以上，在前50关键词中）', size=12)
    add_table(growth_keywords[['Keyword', 'SearchCount_Week1', 'SearchCount_Week2']],
              ['关键词', '本周搜索次数', '上周搜索次数'])

    # 4. 热搜品牌TOP20：从本地读取的 brand 表中筛选符合条件的关键词，品牌名称包括中文和英文
    top_brands = merged_data[merged_data['Keyword'].isin(brand_data['display_name_zh']) | merged_data['Keyword'].isin(
        brand_data['display_name_en'])]
    top_brands = top_brands[['Keyword', 'SearchCount_Week1']].sort_values(by='SearchCount_Week1', ascending=False).head(
        20)
    add_title('热搜品牌TOP20', size=12)
    add_table(top_brands, ['关键词', '搜索次数'])

    # 5. 日仓搜索次数较多的关键词：从日仓关键词中筛选搜索次数大于 50 的关键词
    top_ricang_keywords = merged_data[merged_data['Keyword'].isin(ricang_keywords)]
    top_ricang_keywords = top_ricang_keywords[top_ricang_keywords['SearchCount_Week1'] > 50]
    add_title('日仓搜索次数较多的关键词（搜索次数>50）', size=12)
    add_table(top_ricang_keywords, ['关键词', '搜索次数'])

    # 6. 第三方搜索次数较多的关键词：从第三方关键词中筛选搜索次数大于 50 的关键词
    top_third_party_keywords = merged_data[merged_data['Keyword'].isin(third_party_keywords)]
    top_third_party_keywords = top_third_party_keywords[top_third_party_keywords['SearchCount_Week1'] > 50]
    add_title('第三方搜索次数较多的关键词（搜索次数>50）', size=12)
    add_table(top_third_party_keywords, ['关键词', '搜索次数'])

# 保存文档
save_path = os.path.join('/content/', 'output.docx')
try:
    doc.save(save_path)
    print(f"Word 文档已成功保存为 {save_path}")
except Exception as e:
    print(f"保存文档时出错: {e}")